from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference, PieChart
import os
import time
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.fonts import addMapping
from reportlab.lib.utils import ImageReader
from openpyxl import Workbook
from openpyxl.chart import BarChart, PieChart, Reference
from openpyxl.chart.legend import Legend
from openpyxl.chart.label import DataLabelList
from openpyxl.chart.title import Title
from openpyxl.chart.layout import Layout, ManualLayout
import tempfile
from pathlib import Path
from PIL import Image
import io
import uuid
from pathlib import Path


BASE_DIR = Path(__file__).resolve().parent
TEMP_DIR = BASE_DIR / "temp"
TEMP_DIR.mkdir(exist_ok=True)

def cleanup_temp_files(max_age_hours=1):
    """Удаляет временные файлы старше указанного количества часов"""
    now = time.time()
    for f in TEMP_DIR.glob("*"):
        if f.is_file():
            if now - f.stat().st_mtime > max_age_hours * 3600:
                f.unlink()


def register_cyrillic_fonts():
    """Регистрирует шрифт Times New Roman для поддержки кириллицы"""
    fonts_registered = {'regular': 'Helvetica', 'bold': 'Helvetica-Bold'}

    # Пути к Times New Roman в разных ОС
    possible_paths = [
        # Windows
        "C:\\Windows\\Fonts\\times.ttf",
        "C:\\Windows\\Fonts\\timesbd.ttf",
        "C:\\Windows\\Fonts\\timesi.ttf",
        "C:\\Windows\\Fonts\\timesbi.ttf",
        # Linux
        "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSerif-Regular.ttf",
        # MacOS
        "/Library/Fonts/Times New Roman.ttf",
    ]

    bold_paths = [
        "C:\\Windows\\Fonts\\timesbd.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman_Bold.ttf",
        "/Library/Fonts/Times New Roman Bold.ttf",
    ]

    # Регистрируем обычный шрифт
    regular_path = None
    for path in possible_paths:
        if os.path.exists(path):
            regular_path = path
            break

    # Регистрируем жирный шрифт
    bold_path = None
    for path in bold_paths:
        if os.path.exists(path):
            bold_path = path
            break

    try:
        if regular_path:
            pdfmetrics.registerFont(TTFont('Times-Roman', regular_path))
            fonts_registered['regular'] = 'Times-Roman'
            print(f"✅ Шрифт Times-Roman зарегистрирован: {regular_path}")
        else:
            print("⚠️ Шрифт Times-Roman не найден, пробуем Arial")
            # Пробуем Arial как запасной
            arial_paths = [
                "C:\\Windows\\Fonts\\arial.ttf",
                "/usr/share/fonts/truetype/msttcorefonts/Arial.ttf",
                "/Library/Fonts/Arial.ttf",
            ]
            for path in arial_paths:
                if os.path.exists(path):
                    pdfmetrics.registerFont(TTFont('Times-Roman', path))
                    fonts_registered['regular'] = 'Times-Roman'
                    print(f"✅ Шрифт Arial зарегистрирован как Times-Roman: {path}")
                    break

        if bold_path:
            pdfmetrics.registerFont(TTFont('Times-Bold', bold_path))
            fonts_registered['bold'] = 'Times-Bold'
            print(f"✅ Шрифт Times-Bold зарегистрирован: {bold_path}")
        else:
            # Пробуем Arial Bold
            arial_bold_paths = [
                "C:\\Windows\\Fonts\\arialbd.ttf",
                "/usr/share/fonts/truetype/msttcorefonts/Arial_Bold.ttf",
                "/Library/Fonts/Arial Bold.ttf",
            ]
            for path in arial_bold_paths:
                if os.path.exists(path):
                    pdfmetrics.registerFont(TTFont('Times-Bold', path))
                    fonts_registered['bold'] = 'Times-Bold'
                    print(f"✅ Arial Bold зарегистрирован как Times-Bold: {path}")
                    break

    except Exception as e:
        print(f"❌ Ошибка регистрации шрифтов: {e}")

    return fonts_registered


# Регистрируем шрифты при импорте
FONTS_LOADED = register_cyrillic_fonts()


def generate_contract_docx(contract_data: dict, property_data: dict, tenant_data: dict, owner_data: dict,
                           output_path: str = None):
    """
    Генерирует договор аренды в формате Word.
    Все заголовки черные, шрифт Times New Roman.
    """
    from docx.shared import RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import os

    doc = Document()

    # Настройка полей
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    # Стиль для всего документа - Times New Roman
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)

    # Для заголовков тоже Times New Roman, черный цвет
    heading1 = doc.styles['Heading 1']
    heading1.font.name = 'Times New Roman'
    heading1.font.color.rgb = RGBColor(0, 0, 0)

    heading2 = doc.styles['Heading 2']
    heading2.font.name = 'Times New Roman'
    heading2.font.color.rgb = RGBColor(0, 0, 0)

    # ========== ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ==========
    def detect_gender(full_name: str) -> dict:
        if not full_name:
            return {"gender": "male", "ending": "ий", "aya": "ый"}
        name_parts = full_name.split()
        if len(name_parts) >= 3:
            patronymic = name_parts[2]
            if patronymic.endswith('вна') or patronymic.endswith('чна'):
                return {"gender": "female", "ending": "ая", "aya": "ая"}
        if len(name_parts) >= 2:
            name = name_parts[1] if len(name_parts) > 1 else name_parts[0]
            if name.endswith('а') or name.endswith('я'):
                return {"gender": "female", "ending": "ая", "aya": "ая"}
        return {"gender": "male", "ending": "ий", "aya": "ый"}

    def get_party_declension(party: str, gender: dict) -> str:
        if party == "Арендодатель":
            return "именуем" + ("ая" if gender["gender"] == "female" else "ый")
        elif party == "Арендатор":
            return "именуем" + ("ая" if gender["gender"] == "female" else "ый")
        return party

    # ========== ОПРЕДЕЛЯЕМ СКЛОНЕНИЕ ==========
    tenant_gender = detect_gender(tenant_data.get("name", ""))
    owner_gender = detect_gender(owner_data.get("name", ""))

    # ========== ЗАГОЛОВОК ==========
    title = doc.add_heading('ДОГОВОР АРЕНДЫ', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(18)
    title.runs[0].font.bold = True
    title.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_heading(f'недвижимого имущества № {contract_data.get("number", "___")}', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.bold = True
    subtitle.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()

    # ========== ГОРОД И ДАТА ==========
    #p = doc.add_paragraph()
    #p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    #p.add_run(
    #    f'г. {property_data.get("city", "Москва")}     «{contract_data.get("day", "___")}» {contract_data.get("month", "_____")} {contract_data.get("year", "___")} г.'
    #).font.size = Pt(14)

    #doc.add_paragraph()

    # ========== ПРЕАМБУЛА С ПРАВИЛЬНЫМ СКЛОНЕНИЕМ ==========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.25)

    owner_decl = get_party_declension("Арендодатель", owner_gender)
    tenant_decl = get_party_declension("Арендатор", tenant_gender)

    p.add_run(
        f'{owner_data.get("name", "_______________")}, {owner_decl} в дальнейшем «Арендодатель», '
        f'в лице {owner_data.get("rep", "_______________")}, действующ{owner_gender["ending"]} на основании '
        f'{owner_data.get("basis", "_______________")}, с одной стороны, и '
        f'{tenant_data.get("name", "_______________")}, {tenant_decl} в дальнейшем «Арендатор», '
        f'в лице {tenant_data.get("rep", "_______________")}, действующ{tenant_gender["ending"]} на основании '
        f'{tenant_data.get("basis", "_______________")}, с другой стороны, '
        f'заключили настоящий Договор о нижеследующем:'
    ).font.size = Pt(14)

    doc.add_paragraph()

    # ========== РАЗДЕЛ 1. ПРЕДМЕТ ДОГОВОРА ==========
    heading1 = doc.add_heading('1. ПРЕДМЕТ ДОГОВОРА', level=2)
    for run in heading1.runs:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

    property_type = {
        'apartment': 'квартиру',
        'house': 'дом',
        'commercial': 'нежилое помещение'
    }.get(property_data.get("type"), 'квартиру')

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run(
        f'1.1. Арендодатель обязуется передать Арендатору во временное владение и пользование '
        f'{property_type}, расположенную по адресу: {property_data.get("address", "_______________")}, '
        f'г. {property_data.get("city", "_______________")} (далее – «Помещение»), для использования в целях, '
        f'указанных в п. 1.2 настоящего Договора от «{contract_data.get("day", "___")}» {contract_data.get("month", "_____")} {contract_data.get("year", "___")} г.'
    ).font.size = Pt(14)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run(
        f'1.2. Помещение предоставляется для использования под {contract_data.get("purpose", "проживание")}.'
    ).font.size = Pt(14)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run(
        f'1.3. Характеристики Помещения:\n'
        f'           - общая площадь: {property_data.get("area", "___")} кв. м;\n'
        f'           - количество комнат: {property_data.get("rooms", "___")}.'
    ).font.size = Pt(14)

    doc.add_paragraph()

    # ========== РАЗДЕЛ 2. СРОК АРЕНДЫ ==========
    heading2 = doc.add_heading('2. СРОК АРЕНДЫ', level=2)
    for run in heading2.runs:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run(
        f'2.1. Настоящий Договор заключен сроком на {contract_data.get("duration_months", "___")} месяцев '
        f'и действует с «{contract_data.get("start_day", "___")}» {contract_data.get("start_month", "_____")} '
        f'{contract_data.get("start_year", "___")} г. по «{contract_data.get("end_day", "___")}» '
        f'{contract_data.get("end_month", "_____")} {contract_data.get("end_year", "___")} г.'
    ).font.size = Pt(14)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run(
        f'2.2. Если ни одна из Сторон не заявит о своем намерении прекратить Договор не позднее чем за 30 '
        f'(тридцать) дней до окончания срока его действия, Договор считается продленным на тот же срок '
        f'на тех же условиях.'
    ).font.size = Pt(14)

    doc.add_paragraph()

    # ========== РАЗДЕЛ 3. АРЕНДНАЯ ПЛАТА ==========
    heading3 = doc.add_heading('3. АРЕНДНАЯ ПЛАТА И РАСЧЕТЫ', level=2)
    for run in heading3.runs:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

    interval_text = {
        'month': 'ежемесячно',
        'week': 'еженедельно'
    }.get(property_data.get("interval"), 'ежемесячно')

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run(
        f'3.1. Арендная плата за Помещение устанавливается в размере '
        f'{contract_data.get("monthly_price", "___")} рублей и вносится {interval_text}.'
    ).font.size = Pt(14)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run(
        f'3.2. Арендная плата вносится путем перечисления денежных средств на расчетный счет Арендодателя '
        f'или наличными денежными средствами не позднее {contract_data.get("payment_day", "10")} числа каждого месяца.'
    ).font.size = Pt(14)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run(
        f'3.3. Сумма обеспечительного платежа (депозита) составляет '
        f'{contract_data.get("deposit", "___")} рублей и вносится Арендатором до подписания настоящего Договора. '
        f'Указанная сумма возвращается Арендатору при расторжении Договора при отсутствии задолженности и повреждений имущества.'
    ).font.size = Pt(14)

    doc.add_paragraph()

    # ========== ДАТА ДОГОВОРА В КОНЦЕ ==========
    #p = doc.add_paragraph()
    #p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    #p.add_run(
    #    f'г. {property_data.get("city", "Москва")} «{contract_data.get("day", "___")}» {contract_data.get("month", "_____")} {contract_data.get("year", "___")} г.'
    #).font.size = Pt(14)

    #doc.add_paragraph()

    # ========== ПОДПИСИ СТОРОН С УЧЕТОМ ПОДПИСАНИЯ ==========
    doc.add_heading('ПОДПИСИ СТОРОН', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Таблица для подписей
    table = doc.add_table(rows=5, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(8)

    # Арендодатель
    table.cell(0, 0).text = 'АРЕНДОДАТЕЛЬ:'
    table.cell(1, 0).text = owner_data.get("name", "___________")
    table.cell(2, 0).text = f'Паспорт: {owner_data.get("passport", "___________")}'
    table.cell(3, 0).text = ''  # Пустая строка для отступа

    # ===== ПОДПИСЬ АРЕНДОДАТЕЛЯ =====
    if contract_data.get("owner_signed"):
        # Если подписано, вставляем картинку подписи
        signature_path = Path(__file__).parent / "resources" / "signature.png"
        if signature_path.exists():
            run = table.cell(4, 0).paragraphs[0].add_run()
            run.add_picture(str(signature_path), width=Cm(3), height=Cm(1.5))
            table.cell(4, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            table.cell(4, 0).text = '✓ Подписано'
    else:
        table.cell(4, 0).text = '_______________'

    # Арендатор
    table.cell(0, 1).text = 'АРЕНДАТОР:'
    table.cell(1, 1).text = tenant_data.get("name", "___________")
    table.cell(2, 1).text = f'Паспорт: {tenant_data.get("passport", "___________")}'
    table.cell(3, 1).text = ''  # Пустая строка для отступа

    # ===== ПОДПИСЬ АРЕНДАТОРА =====
    if contract_data.get("tenant_signed"):
        signature_path = Path(__file__).parent / "resources" / "signature.png"
        if signature_path.exists():
            run = table.cell(4, 1).paragraphs[0].add_run()
            run.add_picture(str(signature_path), width=Cm(3), height=Cm(1.5))
            table.cell(4, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            table.cell(4, 1).text = '✓ Подписано'
    else:
        table.cell(4, 1).text = '_______________'

    # ========== СОХРАНЕНИЕ ФАЙЛА ==========
    if output_path:
        doc.save(output_path)
        return output_path
    else:
        temp_dir = Path(__file__).parent / "temp"
        temp_dir.mkdir(exist_ok=True)
        filename = f"contract_{uuid.uuid4().hex}.docx"
        filepath = temp_dir / filename
        doc.save(str(filepath))
        return str(filepath)


def detect_gender(full_name: str) -> dict:
    """
    Определяет пол по окончанию ФИО
    Возвращает словарь с окончаниями для склонения
    """
    if not full_name:
        return {"gender": "male", "ending": "ий", "ie": "ый", "aya": "ый"}

    # Простая проверка по окончанию отчества или имени
    name_parts = full_name.split()

    # Проверяем отчество (обычно заканчивается на "вич" или "вна")
    if len(name_parts) >= 3:
        patronymic = name_parts[2]
        if patronymic.endswith('вна') or patronymic.endswith('чна'):
            return {"gender": "female", "ending": "ая", "ie": "ая", "aya": "ая"}

    # Проверяем имя
    if len(name_parts) >= 2:
        name = name_parts[1] if len(name_parts) > 1 else name_parts[0]
        if name.endswith('а') or name.endswith('я'):
            return {"gender": "female", "ending": "ая", "ie": "ая", "aya": "ая"}

    # По умолчанию - мужской род
    return {"gender": "male", "ending": "ий", "ie": "ый", "aya": "ый"}


def get_party_declension(party: str, gender: dict) -> str:
    """
    Возвращает правильное склонение для стороны договора
    """
    if party == "Арендодатель":
        return "именуем" + gender["ending"]
    elif party == "Арендатор":
        return "именуем" + gender["ending"]
    return party


def generate_agent_stats_excel(agent_id, months, monthly_data, perf_data, status_data):
    """Генерирует Excel-файл со статистикой агента на одном листе"""
    from openpyxl import Workbook
    from openpyxl.chart import BarChart, PieChart, Reference
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    import tempfile
    from pathlib import Path
    import uuid
    import logging

    logger = logging.getLogger(__name__)

    wb = Workbook()
    ws = wb.active
    ws.title = "Статистика агента"

    # Стиль границ
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # ===== ЗАГОЛОВОК =====
    ws.merge_cells('A1:H1')
    title_cell = ws['A1']
    title_cell.value = f"Статистика работы агента за {months} месяцев"
    title_cell.font = Font(size=16, bold=True)
    title_cell.alignment = Alignment(horizontal='center', vertical='center')

    # ===== KPI =====
    ws['A3'] = "КЛЮЧЕВЫЕ ПОКАЗАТЕЛИ"
    ws['A3'].font = Font(size=14, bold=True)

    # Преобразуем perf_data в словарь
    if hasattr(perf_data, '_mapping'):
        perf_dict = dict(perf_data._mapping)
    elif hasattr(perf_data, 'keys'):
        perf_dict = perf_data
    else:
        perf_dict = {}

    kpi_data = [
        ["Общий доход", f"{float(perf_dict.get('total_profit', 0)):,.2f} ₽"],
        ["Средний доход/объект", f"{float(perf_dict.get('avg_profit_per_property', 0)):,.2f} ₽"],
        ["Количество сделок", perf_dict.get('total_deals', 0)],
        ["Загрузка фонда", f"{float(perf_dict.get('occupancy_rate', 0)):.1f}%"],
        ["Обработано заявок", perf_dict.get('processed_applications', 0)],
        ["Среднее время реакции", f"{float(perf_dict.get('avg_response_hours', 0)):.2f} ч"],
        ["Конверсия", f"{float(perf_dict.get('conversion_rate', 0)):.2f}%"]
    ]

    # Заполняем KPI и добавляем границы
    for i, (label, value) in enumerate(kpi_data, start=4):
        ws[f'A{i}'] = label
        ws[f'B{i}'] = value
        ws[f'A{i}'].font = Font(bold=True)

        # Границы для A4:B10
        if 4 <= i <= 10:
            ws[f'A{i}'].border = thin_border
            ws[f'B{i}'].border = thin_border

    # ===== ЕЖЕМЕСЯЧНАЯ СТАТИСТИКА =====
    start_row = 12
    ws[f'A{start_row}'] = "ЕЖЕМЕСЯЧНАЯ СТАТИСТИКА"
    ws[f'A{start_row}'].font = Font(size=14, bold=True)

    # Заголовки
    headers = ["Месяц", "Сделки", "Прибыль (₽)", "Заявки", "Одобрено", "Отклонено"]
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=start_row + 1, column=col)
        cell.value = header
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")
        cell.border = thin_border

    # Преобразуем monthly_data в список словарей
    monthly_list = []
    if monthly_data:
        for row in monthly_data:
            if hasattr(row, '_mapping'):
                monthly_list.append(dict(row._mapping))
            elif isinstance(row, dict):
                monthly_list.append(row)
            elif isinstance(row, (list, tuple)) and len(row) >= 6:
                monthly_list.append({
                    'month': row[0],
                    'deals_count': row[1],
                    'total_profit': row[2],
                    'applications_count': row[3],
                    'approved_count': row[4],
                    'rejected_count': row[5]
                })
            else:
                monthly_list.append(row)

    # Заполняем данные и добавляем границы
    if monthly_list:
        for row_idx, row in enumerate(monthly_list, start=start_row + 2):
            ws.cell(row=row_idx, column=1).value = row.get('month', '')
            ws.cell(row=row_idx, column=2).value = row.get('deals_count', 0)
            ws.cell(row=row_idx, column=3).value = float(row.get('total_profit', 0))
            ws.cell(row=row_idx, column=4).value = row.get('applications_count', 0)
            ws.cell(row=row_idx, column=5).value = row.get('approved_count', 0)
            ws.cell(row=row_idx, column=6).value = row.get('rejected_count', 0)

            # Границы для всех ячеек в диапазоне A(start_row+2):F(последняя строка)
            for col in range(1, 7):
                cell = ws.cell(row=row_idx, column=col)
                cell.border = thin_border

    # ===== СТАТИСТИКА ПО СТАТУСАМ (только approved и rejected) =====
    status_start_row = start_row + len(monthly_list) + 5
    ws[f'A{status_start_row}'] = "СТАТИСТИКА ПО СТАТУСАМ ЗАЯВОК"
    ws[f'A{status_start_row}'].font = Font(size=14, bold=True)

    status_headers = ["Статус", "Количество", "Процент"]
    for col, header in enumerate(status_headers, start=1):
        cell = ws.cell(row=status_start_row + 1, column=col)
        # Переводим статусы на русский
        if header == "Статус":
            cell.value = "Статус"
        else:
            cell.value = header
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")
        cell.border = thin_border

    # Преобразуем status_data в список словарей и фильтруем только approved/rejected
    status_list = []
    total_approved_rejected = 0

    if status_data:
        for row in status_data:
            if hasattr(row, '_mapping'):
                row_dict = dict(row._mapping)
            elif isinstance(row, dict):
                row_dict = row
            elif isinstance(row, (list, tuple)) and len(row) >= 3:
                row_dict = {'status': row[0], 'count': row[1], 'percentage': row[2]}
            else:
                row_dict = row

            # Берем только approved и rejected
            if row_dict.get('status') in ['approved', 'rejected']:
                # Переводим статус на русский
                if row_dict.get('status') == 'approved':
                    row_dict['status'] = 'Одобрено'
                elif row_dict.get('status') == 'rejected':
                    row_dict['status'] = 'Отказано'
                status_list.append(row_dict)
                total_approved_rejected += row_dict.get('count', 0)

    # Пересчитываем проценты для отфильтрованных данных
    if status_list and total_approved_rejected > 0:
        for row in status_list:
            count = row.get('count', 0)
            row['percentage'] = round(count * 100 / total_approved_rejected, 2)

    # Заполняем данные статусов с границами
    if status_list:
        for row_idx, row in enumerate(status_list, start=status_start_row + 2):
            ws.cell(row=row_idx, column=1).value = row.get('status', '')
            ws.cell(row=row_idx, column=2).value = row.get('count', 0)
            ws.cell(row=row_idx, column=3).value = row.get('percentage', 0)

            # Границы для всех ячеек в таблице статусов
            for col in range(1, 4):
                cell = ws.cell(row=row_idx, column=col)
                cell.border = thin_border

    # ===== ДИАГРАММЫ =====
    try:
        if monthly_list and len(monthly_list) > 0:
            # Столбчатая диаграмма
            chart = BarChart()
            chart.title = "Динамика прибыли по месяцам"
            chart.style = 2
            chart.height = 8.55
            chart.width = 12.4

            data_start_row = start_row + 2
            data_end_row = start_row + 1 + len(monthly_list)

            data = Reference(ws, min_col=3, min_row=data_start_row, max_row=data_end_row)
            cats = Reference(ws, min_col=1, min_row=data_start_row, max_row=data_end_row)
            chart.add_data(data, titles_from_data=False)
            chart.set_categories(cats)
            ws.add_chart(chart, f"H{start_row - 10}")

        if status_list and len(status_list) > 0:
            # Круговая диаграмма
            pie = PieChart()
            pie.title = "Соотношение одобрений и отказов"
            pie.height = 6.15
            pie.width = 12.45

            labels = Reference(ws, min_col=1, min_row=status_start_row + 2,
                               max_row=status_start_row + 1 + len(status_list))
            data = Reference(ws, min_col=2, min_row=status_start_row + 2,
                             max_row=status_start_row + 1 + len(status_list))
            pie.add_data(data, titles_from_data=False)
            pie.set_categories(labels)

            ws.add_chart(pie, f"H{start_row + 6}")

    except Exception as e:
        logger.error(f"Ошибка при создании диаграмм: {e}")

    # ===== УСТАНОВКА ШИРИНЫ КОЛОНОК =====
    ws.column_dimensions['A'].width = 44
    ws.column_dimensions['B'].width = 14
    ws.column_dimensions['C'].width = 12
    ws.column_dimensions['D'].width = 7
    ws.column_dimensions['E'].width = 10
    ws.column_dimensions['F'].width = 10
    ws.column_dimensions['G'].width = 3
    for col in range(2, 7):
        ws.column_dimensions[get_column_letter(col)].width = 15

    # Сохраняем
    temp_dir = Path(__file__).parent / "temp"
    temp_dir.mkdir(exist_ok=True)
    filename = f"agent_stats_{agent_id}_{months}_{uuid.uuid4().hex}.xlsx"
    filepath = temp_dir / filename
    wb.save(str(filepath))

    logger.info(f"Excel файл сохранен: {filepath}")
    return str(filepath)


def generate_act_pdf(contract_data: dict, property_data: dict, tenant_data: dict, owner_data: dict,
                     output_path: str = None):
    """
    Генерирует акт приема-передачи в формате PDF с поддержкой кириллицы.
    """
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.utils import ImageReader
    from PIL import Image
    import io

    if output_path is None:
        filename = f"act_{uuid.uuid4().hex}.pdf"
        output_path = str(TEMP_DIR / filename)

    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4

    # Регистрируем шрифты с поддержкой кириллицы
    fonts = register_cyrillic_fonts()
    regular_font = fonts['regular']
    bold_font = fonts['bold']

    # Определяем пол для склонения
    def detect_gender(full_name: str) -> dict:
        if not full_name:
            return {"gender": "male", "ending": "ый"}
        name_parts = full_name.split()
        if len(name_parts) >= 3:
            patronymic = name_parts[2]
            if patronymic.endswith('вна') or patronymic.endswith('чна'):
                return {"gender": "female", "ending": "ая"}
        if len(name_parts) >= 2:
            name = name_parts[1] if len(name_parts) > 1 else name_parts[0]
            if name.endswith('а') or name.endswith('я'):
                return {"gender": "female", "ending": "ая"}
        return {"gender": "male", "ending": "ый"}

    tenant_gender = detect_gender(tenant_data.get("name", ""))
    owner_gender = detect_gender(owner_data.get("name", ""))

    # Заголовок
    c.setFont(bold_font, 16)
    c.setFillColorRGB(0, 0, 0)
    c.drawCentredString(width / 2, height - 40, "АКТ")
    c.setFont(bold_font, 14)
    c.drawCentredString(width / 2, height - 60, "приема-передачи нежилого помещения")

    # Адрес
    #c.setFont(bold_font, 12)
    #c.drawString(50, height - 90, "г. " + property_data.get("city", "Москва"))

    #address = property_data.get("address", "_______________")
    c.setFont(regular_font, 12)
    #c.drawString(50, height - 110, address)

    # Дата
    #date_text = f'«{contract_data.get("day", "___")}» {contract_data.get("month", "_____")} {contract_data.get("year", "___")} г.'
    #c.setFont(regular_font, 12)
    #c.drawString(50, height - 130, date_text)

    # Основной текст
    y = height - 170

    def draw_wrapped_text(text, x, y, max_width, font_name, font_size=12):
        words = text.split()
        lines = []
        current_line = ""

        for word in words:
            test_line = current_line + word + " "
            if c.stringWidth(test_line, font_name, font_size) < max_width:
                current_line = test_line
            else:
                lines.append(current_line)
                current_line = word + " "
        if current_line:
            lines.append(current_line)

        for line in lines:
            c.drawString(x, y, line.strip())
            y -= 18
        return y

    # Преамбула
    owner_decl = "именуемый" if owner_gender["gender"] == "male" else "именуемая"
    tenant_decl = "именуемый" if tenant_gender["gender"] == "male" else "именуемая"

    texts = [
        f'{owner_data.get("name", "_______________")}, {owner_decl} в дальнейшем «Арендодатель», в лице',
        f'{owner_data.get("rep", "_______________")}, действующего на основании {owner_data.get("basis", "_______________")}, передал, а',
        f'{tenant_data.get("name", "_______________")}, {tenant_decl} в дальнейшем «Арендатор», в лице',
        f'{tenant_data.get("rep", "_______________")}, действующего на основании {tenant_data.get("basis", "_______________")}, принял в аренду',
        f'нежилое помещение, расположенное по адресу {property_data.get("address", "_______________")} общей площадью',
        f'{property_data.get("area", "___")} кв. м для использования под {contract_data.get("purpose", "проживание")} согласно договору N',
        f'{contract_data.get("number", "___")} аренды нежилого помещения от',
        f'«{contract_data.get("start_day", "___")}» {contract_data.get("start_month", "_____")} {contract_data.get("start_year", "___")} г.'
    ]

    for text in texts:
        y = draw_wrapped_text(text, 50, y, width - 100, regular_font, 12) - 5

    # Состояние помещения
    y -= 15
    state_text = ('Техническое состояние нежилого помещения удовлетворительное и позволяет использовать его '
                  f'в целях, предусмотренных п. 1.1 указанного Договора аренды')

    y = draw_wrapped_text(state_text, 50, y, width - 100, regular_font, 12) - 10

    # Дата договора в конце
    y -= 20
    c.setFont(regular_font, 12)
    #c.drawString(50, y,
                # f'г. {property_data.get("city", "Москва")} «{contract_data.get("day", "___")}» {contract_data.get("month", "_____")} {contract_data.get("year", "___")} г.')

    # Подписи
    y -= 50
    c.setFont(bold_font, 12)
    c.drawString(50, y, "Арендодатель:")
    c.drawString(width / 2 + 50, y, "Арендатор:")

    y -= 30
    c.setFont(regular_font, 12)

    # ===== ФУНКЦИЯ ДЛЯ ВСТАВКИ ПОДПИСИ =====
    def draw_signature(canvas, x, y, width=100, height=40, is_signed=False):
        if is_signed:
            # Путь к изображению подписи
            signature_path = Path(__file__).parent / "resources" / "signature.png"
            if signature_path.exists():
                try:
                    # Открываем изображение с PIL
                    img = Image.open(signature_path)

                    # Конвертируем в RGBA для обработки прозрачности
                    if img.mode != 'RGBA':
                        img = img.convert('RGBA')

                    # Создаем белый фон (можно убрать, если нужно только прозрачность)
                    # white_bg = Image.new('RGBA', img.size, (255, 255, 255, 255))
                    # combined = Image.alpha_composite(white_bg, img)
                    # combined = combined.convert('RGB')

                    # Сохраняем во временный буфер
                    img_buffer = io.BytesIO()
                    img.save(img_buffer, format='PNG')
                    img_buffer.seek(0)

                    # Рисуем изображение с прозрачностью
                    canvas.drawImage(ImageReader(img_buffer), x, y - height + 20, width=width, height=height,
                                     preserveAspectRatio=True, mask='auto')
                    return True
                except Exception as e:
                    print(f"Ошибка вставки подписи: {e}")
                    canvas.drawString(x, y, '✓ Подписано')
                    return False
            else:
                canvas.drawString(x, y, '✓ Подписано')
                return False
        else:
            canvas.drawString(x, y, '_______________')
            return False

    # Подпись арендодателя
    if contract_data.get("owner_signed"):
        draw_signature(c, 50, y, width=100, height=40, is_signed=True)
    else:
        c.drawString(50, y, '_______________')

    # Подпись арендатора
    if contract_data.get("tenant_signed"):
        draw_signature(c, width / 2 + 50, y, width=100, height=40, is_signed=True)
    else:
        c.drawString(width / 2 + 50, y, '_______________')

    # Паспортные данные
    y -= 25
    c.setFont(regular_font, 10)
    c.drawString(50, y, f'Паспорт: {owner_data.get("passport", "___________")}')
    c.drawString(width / 2 + 50, y, f'Паспорт: {tenant_data.get("passport", "___________")}')

    c.save()
    return output_path